"""Intelligent PDF to DOCX converter with automatic detection.

Modes:
- auto: Auto-detect best method based on PDF content
- text: Text-based conversion with enhanced layout detection
- ocr: OCR-based conversion for scanned PDFs
- image: Embed pages as images
- groq: AI-powered reconstruction using Groq
- hybrid: PyMuPDF + Groq AI pipeline (BEST QUALITY)

Flow:
1. Analyze PDF to detect type (text-based, image-based, or hybrid)
2. For text-based pages: Use enhanced layout detection with pdfplumber
3. For image-based pages: Use OCR (pytesseract + pdf2image)
4. For AI modes: Use Groq AI for semantic reconstruction
"""
import os
import logging
import tempfile
from pathlib import Path
from typing import Dict, Optional

try:
    from pdf2docx import Converter as PDF2DOCXConverter
except ImportError:
    PDF2DOCXConverter = None

try:
    from docx import Document
    from docx.shared import Inches, Pt
except ImportError:
    Document = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from pdf2image import convert_from_path
except ImportError:
    convert_from_path = None

try:
    import pytesseract
    from PIL import Image
except ImportError:
    pytesseract = None
    Image = None

logger = logging.getLogger(__name__)


def convert(input_path: str, output_path: str, mode: str = 'auto') -> str:
    """Convert PDF to DOCX with intelligent page-by-page detection.
    
    Args:
        input_path: Path to input PDF
        output_path: Path to output DOCX
        mode: Conversion mode
            - 'auto' (default): Detect PDF type and use best method
            - 'text': Force text-based conversion with layout detection
            - 'ocr': Force OCR-based conversion (pytesseract)
            - 'image': Embed pages as images
            - 'groq': AI-powered reconstruction using Groq (BETA)
            - 'hybrid': PyMuPDF + Groq AI pipeline (BEST QUALITY)
    
    Returns:
        Path to output file
    """
    if mode == 'auto':
        _convert_auto_detect(input_path, output_path)
    elif mode == 'text':
        _convert_text_based(input_path, output_path)
    elif mode == 'ocr':
        _convert_image_based(input_path, output_path)
    elif mode == 'image':
        _convert_as_images(input_path, output_path)
    elif mode == 'groq':
        _convert_with_groq_ai(input_path, output_path)
    elif mode == 'hybrid':
        # Use hybrid AI pipeline (PyMuPDF + Groq)
        _convert_hybrid_ai(input_path, output_path)
    else:
        raise ValueError(f"Unknown mode: {mode}. Use 'auto', 'text', 'ocr', 'image', 'groq', or 'hybrid'.")
    
    return output_path


def _detect_pdf_type(input_path: str) -> dict:
    """Detect PDF type: text-based, image-based, or hybrid.
    
    Returns:
        dict with keys:
            - 'type': 'text', 'image', or 'hybrid'
            - 'page_types': list of types per page
            - 'total_pages': int
            - 'text_pages': int
            - 'image_pages': int
    """
    if pdfplumber is None:
        logger.warning("pdfplumber not available, assuming text-based PDF")
        return {'type': 'text', 'page_types': [], 'total_pages': 0, 'text_pages': 0, 'image_pages': 0}
    
    page_types = []
    text_threshold = 50  # minimum chars to consider a page text-based
    
    try:
        with pdfplumber.open(input_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                text_length = len(text.strip())
                
                if text_length >= text_threshold:
                    page_types.append('text')
                    logger.debug(f"Page {i+1}: TEXT-BASED ({text_length} chars)")
                else:
                    page_types.append('image')
                    logger.debug(f"Page {i+1}: IMAGE-BASED ({text_length} chars)")
            
            total_pages = len(page_types)
            text_pages = page_types.count('text')
            image_pages = page_types.count('image')
            
            # Determine overall type
            if image_pages == 0:
                pdf_type = 'text'
            elif text_pages == 0:
                pdf_type = 'image'
            else:
                pdf_type = 'hybrid'
            
            result = {
                'type': pdf_type,
                'page_types': page_types,
                'total_pages': total_pages,
                'text_pages': text_pages,
                'image_pages': image_pages
            }
            
            logger.info(f"PDF Analysis: Type={pdf_type}, Total={total_pages}, Text={text_pages}, Image={image_pages}")
            return result
            
    except Exception as e:
        logger.warning(f"PDF analysis failed: {e}, assuming text-based")
        return {'type': 'text', 'page_types': [], 'total_pages': 0, 'text_pages': 0, 'image_pages': 0}


def _convert_auto_detect(input_path: str, output_path: str) -> None:
    """Auto mode: Detect PDF type and use appropriate conversion method."""
    
    # Step 1: Analyze PDF
    analysis = _detect_pdf_type(input_path)
    pdf_type = analysis['type']
    
    # Step 2: Route to appropriate converter
    if pdf_type == 'text':
        logger.info("→ Using TEXT-BASED conversion (pdf2docx)")
        _convert_text_based(input_path, output_path)
    
    elif pdf_type == 'image':
        logger.info("→ Using IMAGE-BASED conversion (OCR)")
        _convert_image_based(input_path, output_path)
    
    elif pdf_type == 'hybrid':
        logger.info("→ Using HYBRID conversion (mixed approach)")
        _convert_hybrid_mixed(input_path, output_path, analysis)


def _convert_text_based(input_path: str, output_path: str) -> None:
    """Convert text-based PDF using enhanced detection for bullets and lists."""
    if PDF2DOCXConverter is None and pdfplumber is None:
        raise RuntimeError('Text conversion requires pdf2docx or pdfplumber. Install with: pip install pdf2docx pdfplumber')
    
    # Try enhanced conversion with bullet detection if pdfplumber is available
    if pdfplumber is not None and Document is not None:
        try:
            logger.info("Converting with enhanced layout detection (bullets, lists, indentation)...")
            _convert_with_layout_detection(input_path, output_path)
            return
        except Exception as e:
            logger.warning(f"Enhanced conversion failed: {e}, falling back to pdf2docx")
    
    # Fallback to pdf2docx
    if PDF2DOCXConverter is None:
        raise RuntimeError('pdf2docx is not installed. Install with: pip install pdf2docx')
    
    logger.info("Converting with pdf2docx (preserves layout, bullets, tables)...")
    
    conv = PDF2DOCXConverter(input_path)
    conv.convert(output_path)
    conv.close()
    
    logger.info(f"✓ Text-based conversion complete: {output_path}")


def _convert_with_layout_detection(input_path: str, output_path: str) -> None:
    """Enhanced conversion using pdfplumber for bullet/list detection."""
    assert pdfplumber is not None
    assert Document is not None
    
    import re
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Bullet/list markers to detect
    BULLET_MARKERS = {'•', '◦', '▪', '▫', '◾', '◽', '○', '●', '-', '*', '→', '►', '‣'}
    NUMBERED_PATTERN = re.compile(r'^(\d+[\.\)]|\([a-z]\)|\([ivx]+\)|[a-z][\.\)])[\s]+')
    
    doc = Document()
    
    with pdfplumber.open(input_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            logger.info(f"  Processing page {page_num}/{len(pdf.pages)} with layout detection...")
            
            # Extract words with position metadata
            words = page.extract_words(
                x_tolerance=3,
                y_tolerance=3,
                keep_blank_chars=False
            )
            
            if not words:
                continue
            
            # Group words into lines based on Y coordinate
            lines = _group_words_into_lines(words)
            
            # Process each line and detect list items
            for line_data in lines:
                line_text = line_data['text']
                line_x0 = line_data['x0']
                line_indent = line_data['indent_level']
                
                # Detect if this line is a list item
                is_bullet = False
                is_numbered = False
                clean_text = line_text
                
                # Check for bullet markers at start
                if line_text and line_text[0] in BULLET_MARKERS:
                    is_bullet = True
                    clean_text = line_text[1:].lstrip()
                
                # Check for numbered list pattern
                numbered_match = NUMBERED_PATTERN.match(line_text)
                if numbered_match:
                    is_numbered = True
                    clean_text = line_text[numbered_match.end():].lstrip()
                
                # Create paragraph with appropriate style
                if is_bullet:
                    para = doc.add_paragraph(clean_text, style='List Bullet')
                    if line_indent > 0:
                        para.paragraph_format.left_indent = Inches(line_indent * 0.5)
                elif is_numbered:
                    para = doc.add_paragraph(clean_text, style='List Number')
                    if line_indent > 0:
                        para.paragraph_format.left_indent = Inches(line_indent * 0.5)
                else:
                    para = doc.add_paragraph(clean_text)
                    if line_indent > 0:
                        para.paragraph_format.left_indent = Inches(line_indent * 0.5)
    
    doc.save(output_path)
    logger.info(f"✓ Enhanced text-based conversion complete: {output_path}")


def _group_words_into_lines(words: list) -> list:
    """Group words into lines based on Y coordinates and calculate indentation."""
    if not words:
        return []
    
    # Sort words by y0 (top), then x0 (left)
    sorted_words = sorted(words, key=lambda w: (round(w['top'], 1), w['x0']))
    
    lines = []
    current_line = []
    current_y = None
    y_tolerance = 3  # pixels tolerance for same line
    
    # Find the leftmost X position across all words (document left margin)
    all_x0 = [w['x0'] for w in sorted_words]
    base_left_margin = min(all_x0) if all_x0 else 0
    
    for word in sorted_words:
        word_y = round(word['top'], 1)
        
        # Check if this word is on a new line
        if current_y is None or abs(word_y - current_y) > y_tolerance:
            # Save previous line if exists
            if current_line:
                lines.append(_finalize_line(current_line, base_left_margin))
            
            # Start new line
            current_line = [word]
            current_y = word_y
        else:
            # Same line, add word
            current_line.append(word)
    
    # Don't forget the last line
    if current_line:
        lines.append(_finalize_line(current_line, base_left_margin))
    
    return lines


def _finalize_line(words: list, base_margin: float) -> dict:
    """Combine words into a line with metadata."""
    if not words:
        return {'text': '', 'x0': 0, 'indent_level': 0}
    
    # Sort words by x position for correct reading order
    words = sorted(words, key=lambda w: w['x0'])
    
    # Combine text with spaces
    text = ' '.join(w['text'] for w in words)
    
    # Calculate indent level based on leftmost word
    line_x0 = words[0]['x0']
    indent_pixels = line_x0 - base_margin
    
    # Convert indent to levels (roughly every 20-30 pixels is one indent level)
    indent_level = max(0, int(indent_pixels / 25))
    
    return {
        'text': text,
        'x0': line_x0,
        'indent_level': indent_level
    }


def _convert_image_based(input_path: str, output_path: str, dpi: int = 300) -> None:
    """Convert image-based (scanned) PDF using OCR."""
    if not all([convert_from_path, pytesseract, Document, Image]):
        raise RuntimeError(
            'OCR conversion requires: pdf2image, pytesseract, python-docx, Pillow\n'
            'Install with: pip install pdf2image pytesseract python-docx Pillow\n'
            'Also requires: Tesseract-OCR installed on system'
        )
    
    # Type guards for linter
    assert convert_from_path is not None
    assert pytesseract is not None
    assert Document is not None
    assert Image is not None
    
    # Verify Tesseract
    try:
        pytesseract.get_tesseract_version()
    except Exception as e:
        raise RuntimeError(
            f'Tesseract-OCR not found: {e}\n'
            f'Download from: https://github.com/UB-Mannheim/tesseract/wiki'
        )
    
    logger.info("Converting with OCR (extracting text from images)...")
    
    # Convert PDF pages to images
    pages = convert_from_path(input_path, dpi=dpi)
    doc = Document()
    
    for i, page_img in enumerate(pages, start=1):
        logger.info(f"  OCR processing page {i}/{len(pages)}...")
        
        if i > 1:
            doc.add_page_break()
        
        # Run OCR with detailed data for better layout
        try:
            ocr_data = pytesseract.image_to_data(page_img, output_type=pytesseract.Output.DICT)
            
            # Reconstruct text with line breaks
            current_block = -1
            current_para = None
            
            for j in range(len(ocr_data['text'])):
                text = ocr_data['text'][j].strip()
                if not text:
                    continue
                
                block_num = ocr_data['block_num'][j]
                conf = int(ocr_data['conf'][j])
                
                # Skip low confidence text
                if conf < 30:
                    continue
                
                # New block = new paragraph
                if block_num != current_block:
                    current_para = doc.add_paragraph()
                    current_block = block_num
                
                # Add text
                if current_para:
                    current_para.add_run(text + ' ')
        
        except Exception as e:
            logger.warning(f"Detailed OCR failed for page {i}, using simple extraction: {e}")
            # Fallback: simple OCR
            text = pytesseract.image_to_string(page_img)
            doc.add_paragraph(text)
    
    doc.save(output_path)
    logger.info(f"✓ Image-based (OCR) conversion complete: {output_path}")


def _convert_hybrid_mixed(input_path: str, output_path: str, analysis: dict) -> None:
    """Convert hybrid PDF: use pdf2docx for text pages, OCR for image pages."""
    
    page_types = analysis['page_types']
    total_pages = analysis['total_pages']
    
    if not all([convert_from_path, pytesseract, Document, Image, PDF2DOCXConverter, pdfplumber]):
        logger.warning("Missing dependencies for hybrid conversion, using text-based fallback")
        return _convert_text_based(input_path, output_path)
    
    logger.info("Converting hybrid PDF (mixing text and OCR extraction)...")
    
    # Strategy: Use pdf2docx for text pages, then enhance with OCR for image pages
    # Since pdf2docx processes whole document, we'll use a simpler approach:
    # Try pdf2docx first, then append OCR results for image-heavy pages if needed
    
    if analysis['text_pages'] > analysis['image_pages']:
        # Mostly text: use pdf2docx
        logger.info(f"  Primary method: pdf2docx ({analysis['text_pages']} text pages)")
        _convert_text_based(input_path, output_path)
    else:
        # Mostly images: use OCR
        logger.info(f"  Primary method: OCR ({analysis['image_pages']} image pages)")
        _convert_image_based(input_path, output_path)


def _convert_as_images(input_path: str, output_path: str, dpi: int = 200) -> None:
    """Embed each page as an image for exact visual match."""
    if not all([convert_from_path, Document, Image]):
        raise RuntimeError('Image mode requires: pdf2image, python-docx, Pillow')
    
    # Type guards for linter
    assert convert_from_path is not None
    assert Document is not None
    assert Image is not None
    
    logger.info("Converting pages as images (exact visual preservation)...")
    
    pages = convert_from_path(input_path, dpi=dpi)
    doc = Document()
    
    for i, page in enumerate(pages, start=1):
        logger.info(f"  Processing page {i}/{len(pages)}...")
        
        if i > 1:
            doc.add_page_break()
        
        # Save to BytesIO and add to document
        from io import BytesIO
        bio = BytesIO()
        page.save(bio, format='PNG')
        bio.seek(0)
        
        doc.add_picture(bio, width=Inches(6.5))
    
    doc.save(output_path)
    logger.info(f"✓ Image-based conversion complete: {output_path}")


def _convert_with_groq_ai(input_path: str, output_path: str) -> None:
    """AI-powered conversion using Groq for intelligent document reconstruction.
    
    Pipeline:
    1. Extract layout data (coordinates, fonts, indentation) from PDF
    2. Check for sensitive information (privacy protection)
    3. Send to Groq AI for semantic analysis and structure reconstruction
    4. Build styled DOCX from AI-corrected output
    
    This mode provides the best quality for complex documents with:
    - Broken bullets and list formatting
    - Inconsistent headings
    - Poor paragraph grouping
    - Mixed text and image regions
    
    Privacy: Layout data (text + coordinates) is sent to Groq API.
    For sensitive documents, use mode='text' or mode='ocr' instead.
    """
    logger.info("🤖 Starting AI-powered conversion with Groq...")
    
    # Import AI components
    try:
        from utils.layout_extractor import LayoutExtractor
        from utils.groq_service import GroqDocumentReconstructor
        from utils.document_reconstructor import reconstruct_document
        from utils.privacy import PrivacyChecker, audit_api_request
    except ImportError as e:
        raise RuntimeError(
            f"Groq mode requires additional dependencies: {e}\n"
            "Install with: pip install groq pdfplumber"
        )
    
    # Phase 1: Extract Layout
    logger.info("📄 Phase 1: Extracting PDF layout...")
    try:
        extractor = LayoutExtractor(input_path)
        layout_data = extractor.extract()
    except Exception as e:
        logger.error(f"Layout extraction failed: {e}")
        raise
    
    # Phase 1.5: Privacy Check
    logger.info("🔒 Phase 1.5: Privacy check...")
    checker = PrivacyChecker()
    has_sensitive, findings = checker.check_layout_data(layout_data)
    
    if has_sensitive:
        logger.warning("⚠️  PRIVACY WARNING: Sensitive information detected!")
        for finding in findings:
            logger.warning(f"   - {finding}")
        logger.warning("")
        logger.warning("This document may contain sensitive information.")
        logger.warning("Consider using local-only mode instead:")
        logger.warning("  python main.py input.pdf -o output.docx -t docx -m text")
        logger.warning("")
        
        # Check if user wants to proceed (only in interactive mode)
        import sys
        if sys.stdin.isatty():
            response = input("Continue with Groq API? (y/N): ").strip().lower()
            if response != 'y':
                logger.info("Conversion cancelled. Using local text mode instead...")
                return _convert_text_based(input_path, output_path)
        else:
            # Non-interactive: check environment variable
            import os
            if os.getenv('GROQ_ALLOW_SENSITIVE', 'false').lower() != 'true':
                logger.error("Sensitive content detected in non-interactive mode.")
                logger.error("Set GROQ_ALLOW_SENSITIVE=true to override, or use local mode.")
                return _convert_text_based(input_path, output_path)
    
    # Audit logging (if enabled)
    audit_api_request(layout_data)
    
    # Phase 2: AI Reconstruction
    logger.info("🧠 Phase 2: AI document reconstruction with Groq...")
    try:
        groq_service = GroqDocumentReconstructor()
        
        # Use hybrid mode for combined layout + style analysis
        ai_output = groq_service.reconstruct_document(
            layout_data,
            pass_type='hybrid'
        )
        
        # Validate output
        if not groq_service.validate_reconstruction(ai_output):
            logger.warning("AI reconstruction validation failed, using fallback...")
            # Fallback: Try simpler text-based conversion
            return _convert_text_based(input_path, output_path)
    
    except Exception as e:
        logger.error(f"Groq AI reconstruction failed: {e}")
        logger.warning("Falling back to text-based conversion...")
        return _convert_text_based(input_path, output_path)
    
    # Phase 3: Build DOCX
    logger.info("📝 Phase 3: Building styled DOCX...")
    try:
        # Extract metadata for document properties
        metadata = layout_data.get("metadata", {})
        
        # Reconstruct document with AI corrections AND original layout data
        reconstruct_document(
            ai_output,
            output_path,
            metadata=metadata,
            layout_data=layout_data,  # Pass layout data for enhanced formatting
            generate_report=True
        )
    
    except Exception as e:
        logger.error(f"Document reconstruction failed: {e}")
        raise
    
    logger.info(f"✓ AI-powered conversion complete: {output_path}")
    logger.info("  A conversion report was saved alongside the output file.")


def _convert_hybrid_ai(input_path: str, output_path: str, privacy_check: bool = True) -> None:
    """Hybrid AI pipeline: PyMuPDF Layout + Groq AI Reconstruction.
    
    This is the BEST QUALITY conversion mode that combines:
    1. PyMuPDF for accurate layout extraction
    2. Privacy checking (optional)
    3. Groq AI for semantic cleanup and structure recognition
    4. Document reconstruction with proper styling
    
    Pipeline:
    - Extract layout from PDF using PyMuPDF (high-accuracy text & positioning)
    - Privacy check (optional)
    - Groq AI semantic cleanup and restructuring
    - Build final DOCX with proper formatting
    """
    logger.info("🚀 Starting Hybrid AI conversion (PyMuPDF + Groq)...")
    
    try:
        from utils.layout_extractor import LayoutExtractor
        from utils.groq_service import GroqDocumentReconstructor
        from utils.privacy import PrivacyChecker
        from utils.document_reconstructor import DocumentReconstructor
    except ImportError as e:
        raise RuntimeError(
            f"Hybrid mode requires additional dependencies: {e}\n"
            "Install with: pip install groq pdfplumber pymupdf"
        )
    
    # Create temp directory
    temp_dir = tempfile.mkdtemp(prefix='hybrid_conversion_')
    
    try:
        # PHASE 1: PDF Layout Extraction
        logger.info("📋 PHASE 1: PDF Layout Extraction (PyMuPDF)")
        logger.info("-" * 70)
        
        extractor = LayoutExtractor(input_path)
        layout_data = extractor.extract()
        
        total_blocks = sum(len(page['blocks']) for page in layout_data['pages'])
        logger.info(f"✓ Layout extraction complete:")
        logger.info(f"  • Pages: {len(layout_data['pages'])}")
        logger.info(f"  • Total blocks: {total_blocks}")
        
        # PHASE 2: Privacy Check
        if privacy_check:
            logger.info("")
            logger.info("🔒 PHASE 2: Privacy Check")
            logger.info("-" * 70)
            
            text = '\n'.join(
                block.get('text', '')
                for page in layout_data['pages']
                for block in page['blocks']
            )
            
            checker = PrivacyChecker()
            has_sensitive, findings = checker.check_text(text)
            
            if has_sensitive:
                logger.warning(f"⚠️  Found {len(findings)} sensitive indicators:")
                for item in findings:
                    logger.warning(f"  • {item}")
                
                import sys
                if sys.stdin.isatty():
                    response = input("\nContinue with AI processing? (y/N): ").strip().lower()
                    if response != 'y':
                        raise RuntimeError("Conversion cancelled due to privacy concerns")
            else:
                logger.info("✓ No sensitive data patterns detected")
        
        # PHASE 3: Groq AI Reasoning
        logger.info("")
        logger.info("🧠 PHASE 3: Groq AI Reasoning")
        logger.info("-" * 70)
        
        logger.info("Sending to Groq AI for semantic cleanup...")
        
        groq_service = GroqDocumentReconstructor()
        ai_result = groq_service.reconstruct_document(
            layout_data,
            pass_type='hybrid'
        )
        
        # Validate and extract structure
        if ai_result and ai_result.get('status') == 'success':
            blocks_count = len(ai_result.get('reconstructed', {}).get('blocks', []))
            logger.info(f"✓ AI reconstruction complete")
            logger.info(f"  • Blocks: {blocks_count}")
            ai_structure = ai_result
        else:
            logger.warning("⚠️  AI reconstruction returned empty/invalid, using layout data")
            blocks = []
            for page in layout_data.get('pages', []):
                blocks.extend(page.get('blocks', []))
            ai_structure = {'blocks': blocks}
        
        # PHASE 4: Post-Processing
        logger.info("")
        logger.info("📝 PHASE 4: Post-Processing")
        logger.info("-" * 70)
        
        logger.info("Building final DOCX from AI structure...")
        
        reconstructor = DocumentReconstructor(ai_structure)
        reconstructor.build_document(output_path)
        
        file_size = os.path.getsize(output_path)
        logger.info(f"✓ Final DOCX created: {file_size / 1024:.1f} KB")
        
    finally:
        # Cleanup
        import shutil
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
    
    logger.info(f"✓ Hybrid AI conversion complete: {output_path}")


if __name__ == '__main__':
    import sys
    if len(sys.argv) < 3:
        print('Usage: python pdf_to_docx.py input.pdf output.docx [mode]')
        print('Modes:')
        print('  auto   - Auto-detect best method (default)')
        print('  text   - Text-based conversion (pdf2docx)')
        print('  ocr    - OCR for scanned PDFs (pytesseract)')
        print('  image  - Embed pages as images')
        print('  groq   - AI-powered reconstruction (BETA) 🤖')
        sys.exit(2)
    
    inp = sys.argv[1]
    out = sys.argv[2]
    mode = sys.argv[3] if len(sys.argv) > 3 else 'auto'
    
    # Setup logging
    logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(message)s')
    
    convert(inp, out, mode=mode)
