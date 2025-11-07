"""Document reconstructor: Convert AI output to styled DOCX.

This module takes the AI-reconstructed document structure and builds
a properly formatted Word document with correct styles, indentation,
and list formatting.
"""
import logging
from pathlib import Path
from typing import Dict, List, Any, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    Document = None

logger = logging.getLogger(__name__)


# Style mapping from AI tags to Word styles
STYLE_MAP = {
    "Title": "Title",
    "Heading 1": "Heading 1",
    "Heading 2": "Heading 2",
    "Heading 3": "Heading 3",
    "Heading 4": "Heading 4",
    "Heading 5": "Heading 5",
    "List Bullet": "List Bullet",
    "List Number": "List Number",
    "Normal": "Normal",
    "Quote": "Quote",
    "Caption": "Caption",
    "Intense Quote": "Intense Quote",
    "Subtitle": "Subtitle"
}


class DocumentReconstructor:
    """Rebuild DOCX from AI-reconstructed structure."""
    
    def __init__(self, reconstructed_data: Dict[str, Any], layout_data: Optional[Dict[str, Any]] = None):
        """Initialize reconstructor.
        
        Args:
            reconstructed_data: AI-reconstructed document structure
            layout_data: Original layout data with font/color information
        """
        if Document is None:
            raise RuntimeError(
                "python-docx is required. Install with: pip install python-docx"
            )
        
        self.data = reconstructed_data
        self.layout_data = layout_data or {}
        self.doc = Document()
        
        # Track list state for proper numbering continuation
        self.list_state = {
            "bullet_level": 0,
            "number_level": 0,
            "last_style": None
        }
        
        # Build font/style lookup from original layout
        self._build_style_lookup()
        
        logger.info("Document reconstructor initialized")
    
    def _build_style_lookup(self) -> None:
        """Build lookup table for original font/style information."""
        self.style_lookup = {}
        
        if not self.layout_data or "pages" not in self.layout_data:
            return
        
        # Index original blocks by text for quick lookup
        for page in self.layout_data.get("pages", []):
            for block in page.get("blocks", []):
                text_key = block.get("text", "").strip()[:100]  # First 100 chars as key
                if text_key:
                    self.style_lookup[text_key] = {
                        "font": block.get("font", "Calibri"),
                        "size": block.get("size", 11.0),
                        "is_bold": "bold" in block.get("font", "").lower(),
                        "is_italic": "italic" in block.get("font", "").lower(),
                        "indent_level": block.get("indent_level", 0)
                    }
    
    def build_document(self, output_path: str) -> None:
        """Build and save the DOCX document.
        
        Args:
            output_path: Path to save the output DOCX
        """
        logger.info("Building DOCX from AI-reconstructed structure...")
        
        # Get blocks from reconstructed data
        blocks = self._extract_blocks()
        
        if not blocks:
            logger.warning("No blocks found in reconstructed data")
            # Add a placeholder paragraph
            self.doc.add_paragraph("(Empty document)")
        else:
            # Process each block
            for i, block in enumerate(blocks):
                self._add_block_to_document(block, i)
        
        # Apply document-level formatting
        self._apply_document_settings()
        
        # Save
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        self.doc.save(str(output_file))
        logger.info(f"✓ DOCX saved to: {output_file}")
    
    def _extract_blocks(self) -> List[Dict[str, Any]]:
        """Extract blocks from reconstructed data.
        
        Returns:
            List of block dicts
        """
        # Handle different possible structures
        if self.data.get("status") == "success":
            reconstructed = self.data.get("reconstructed", {})
            blocks = reconstructed.get("blocks", [])
        elif "blocks" in self.data:
            blocks = self.data["blocks"]
        else:
            # Fallback: treat as raw text
            logger.warning("Unrecognized data structure, treating as plain text")
            raw_text = self.data.get("raw_text", "")
            if raw_text:
                blocks = [{"text": raw_text, "style": "Normal"}]
            else:
                blocks = []
        
        return blocks
    
    def _add_block_to_document(self, block: Dict[str, Any], index: int) -> None:
        """Add a single block to the document with enhanced formatting.
        
        Args:
            block: Block dict with text and style
            index: Block index for logging
        """
        text = block.get("text", "").strip()
        style_name = block.get("style", "Normal")
        level = block.get("level", 0)
        
        if not text:
            logger.debug(f"Block {index}: Skipping empty block")
            return
        
        # Map AI style to Word style
        word_style = STYLE_MAP.get(style_name, "Normal")
        
        logger.debug(f"Block {index}: Adding '{text[:50]}...' as {word_style}")
        
        try:
            # Create paragraph with appropriate style
            if word_style in ["List Bullet", "List Number"]:
                para = self._add_list_item(text, word_style, level)
            else:
                para = self.doc.add_paragraph(text, style=word_style)
                
                # Apply indentation if specified
                if level > 0:
                    para.paragraph_format.left_indent = Inches(level * 0.5)
            
            # Apply enhanced formatting from original layout
            self._apply_enhanced_formatting(para, text, style_name)
        
        except Exception as e:
            logger.warning(f"Failed to add block {index} with style {word_style}: {e}")
            # Fallback to normal paragraph
            self.doc.add_paragraph(text)
    
    def _apply_enhanced_formatting(self, para: Any, text: str, style_name: str) -> None:
        """Apply enhanced formatting based on original layout data.
        
        Args:
            para: Paragraph object
            text: Text content
            style_name: Style name from AI
        """
        # Look up original formatting
        text_key = text[:100]
        original_style = self.style_lookup.get(text_key)
        
        if not original_style:
            # Apply default formatting improvements
            self._apply_default_formatting(para, style_name)
            return
        
        # Apply font formatting to all runs in paragraph
        for run in para.runs:
            # Font name
            font_name = original_style.get("font", "Calibri")
            # Simplify font name (remove -Bold, -Italic suffixes)
            base_font = font_name.split('-')[0] if '-' in font_name else font_name
            run.font.name = base_font
            
            # Font size (convert to Pt)
            font_size = original_style.get("size", 11.0)
            run.font.size = Pt(font_size)
            
            # Bold/Italic
            if original_style.get("is_bold"):
                run.font.bold = True
            if original_style.get("is_italic"):
                run.font.italic = True
        
        # Paragraph spacing
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.space_before = Pt(0)
        
        # Line spacing
        para.paragraph_format.line_spacing = 1.15
    
    def _apply_default_formatting(self, para: Any, style_name: str) -> None:
        """Apply default formatting improvements when original style not found.
        
        Args:
            para: Paragraph object
            style_name: Style name from AI
        """
        # Set default font and spacing
        for run in para.runs:
            run.font.name = 'Calibri'
            
            # Size based on style
            if style_name == "Title":
                run.font.size = Pt(26)
                run.font.bold = True
            elif style_name == "Heading 1":
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(31, 78, 120)  # Dark blue
            elif style_name == "Heading 2":
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(79, 129, 189)  # Blue
            elif style_name == "Heading 3":
                run.font.size = Pt(12)
                run.font.bold = True
            elif style_name in ["List Bullet", "List Number"]:
                run.font.size = Pt(11)
            else:
                run.font.size = Pt(11)
        
        # Paragraph spacing
        if style_name.startswith("Heading"):
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
        else:
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.space_before = Pt(0)
        
        # Line spacing
        para.paragraph_format.line_spacing = 1.15
    
    def _add_list_item(self, text: str, style: str, level: int) -> Any:
        """Add a list item with proper indentation and formatting.
        
        Args:
            text: Item text
            style: List Bullet or List Number
            level: Indent level (0-based)
        
        Returns:
            Paragraph object
        """
        # Create list paragraph
        para = self.doc.add_paragraph(text, style=style)
        
        # Apply indentation for nested lists
        if level > 0:
            para.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
            para.paragraph_format.first_line_indent = Inches(-0.25)
        else:
            # Standard list indentation
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.25)
        
        # Reduce spacing between list items
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.space_before = Pt(0)
        
        # Check if we need to continue or restart numbering
        current_style = "bullet" if style == "List Bullet" else "number"
        
        if self.list_state["last_style"] != current_style:
            # Style changed, this is a new list
            self.list_state["last_style"] = current_style
        
        return para
    
    def _apply_document_settings(self) -> None:
        """Apply document-level settings (margins, spacing, etc.)."""
        
        # Set standard margins (1 inch = 914400 EMUs)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        logger.debug("Applied document settings (margins, etc.)")
    
    def add_metadata(self, metadata: Dict[str, Any]) -> None:
        """Add metadata to the document.
        
        Args:
            metadata: Metadata dict (title, author, etc.)
        """
        core_props = self.doc.core_properties
        
        if "title" in metadata and metadata["title"]:
            core_props.title = metadata["title"]
        
        if "author" in metadata and metadata["author"]:
            core_props.author = metadata["author"]
        
        if "subject" in metadata and metadata["subject"]:
            core_props.subject = metadata["subject"]
        
        logger.debug(f"Added metadata: {metadata}")
    
    def generate_report(self, output_path: str) -> None:
        """Generate a conversion report.
        
        Args:
            output_path: Path to save the report
        """
        report_lines = [
            "=== AI-Powered PDF to DOCX Conversion Report ===",
            "",
            f"Status: {self.data.get('status', 'unknown')}",
            ""
        ]
        
        # Count blocks by style
        blocks = self._extract_blocks()
        style_counts = {}
        
        for block in blocks:
            style = block.get("style", "Normal")
            style_counts[style] = style_counts.get(style, 0) + 1
        
        report_lines.append("Block Statistics:")
        report_lines.append(f"  Total blocks: {len(blocks)}")
        
        for style, count in sorted(style_counts.items()):
            report_lines.append(f"  {style}: {count}")
        
        report_lines.append("")
        
        # AI reconstruction details
        if self.data.get("status") == "success":
            report_lines.append("AI Reconstruction: SUCCESS")
            
            reconstructed = self.data.get("reconstructed", {})
            if "ai_notes" in reconstructed:
                report_lines.append(f"  Notes: {reconstructed['ai_notes']}")
        else:
            report_lines.append("AI Reconstruction: FAILED or INCOMPLETE")
            if "raw_text" in self.data:
                report_lines.append("  Fallback: Used raw text output")
        
        report_lines.append("")
        report_lines.append("=== End of Report ===")
        
        # Save report
        report_path = Path(output_path)
        report_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(report_lines))
        
        logger.info(f"✓ Conversion report saved to: {report_path}")


def reconstruct_document(
    ai_output: Dict[str, Any],
    output_path: str,
    metadata: Optional[Dict[str, Any]] = None,
    layout_data: Optional[Dict[str, Any]] = None,
    generate_report: bool = False
) -> None:
    """High-level function to reconstruct DOCX from AI output.
    
    Args:
        ai_output: AI-reconstructed document data
        output_path: Path to save output DOCX
        metadata: Optional metadata to add to document
        layout_data: Original layout data for enhanced formatting
        generate_report: Whether to generate a conversion report
    """
    reconstructor = DocumentReconstructor(ai_output, layout_data)
    
    # Add metadata if provided
    if metadata:
        reconstructor.add_metadata(metadata)
    
    # Build document
    reconstructor.build_document(output_path)
    
    # Generate report if requested
    if generate_report:
        report_path = str(Path(output_path).with_suffix('.report.txt'))
        reconstructor.generate_report(report_path)


class LayoutDebugger:
    """Debug tool to visualize layout extraction results."""
    
    @staticmethod
    def print_layout_summary(layout_data: Dict[str, Any]) -> None:
        """Print a human-readable summary of layout data.
        
        Args:
            layout_data: Layout data from LayoutExtractor
        """
        print("\n" + "="*60)
        print("LAYOUT EXTRACTION SUMMARY")
        print("="*60)
        
        metadata = layout_data.get("metadata", {})
        print(f"\nFile: {metadata.get('filename', 'unknown')}")
        print(f"Size: {metadata.get('file_size_mb', 0)} MB")
        print(f"Total pages: {layout_data.get('total_pages', 0)}")
        
        pages = layout_data.get("pages", [])
        
        for page in pages[:3]:  # Show first 3 pages
            page_num = page.get("page", 0)
            blocks = page.get("blocks", [])
            
            print(f"\n--- Page {page_num} ({len(blocks)} blocks) ---")
            
            for i, block in enumerate(blocks[:10]):  # Show first 10 blocks
                text = block.get("text", "")[:60]
                style_hint = ""
                
                if block.get("is_potential_heading"):
                    style_hint = " [HEADING?]"
                elif block.get("list_type"):
                    style_hint = f" [{block['list_type'].upper()} LIST]"
                
                indent = "  " * block.get("indent_level", 0)
                
                print(f"  {indent}Block {i}: {text}{style_hint}")
        
        if len(pages) > 3:
            print(f"\n... ({len(pages) - 3} more pages)")
        
        print("\n" + "="*60 + "\n")
    
    @staticmethod
    def save_layout_html(layout_data: Dict[str, Any], output_path: str) -> None:
        """Save layout visualization as HTML.
        
        Args:
            layout_data: Layout data from LayoutExtractor
            output_path: Path to save HTML file
        """
        html_lines = [
            "<!DOCTYPE html>",
            "<html>",
            "<head>",
            "<meta charset='utf-8'>",
            "<title>Layout Extraction Visualization</title>",
            "<style>",
            "body { font-family: Arial, sans-serif; margin: 20px; }",
            ".page { border: 1px solid #ccc; margin: 20px 0; padding: 20px; }",
            ".block { margin: 10px 0; padding: 8px; border-left: 3px solid #007bff; }",
            ".heading { background: #fff3cd; border-left-color: #856404; font-weight: bold; }",
            ".list { background: #d1ecf1; border-left-color: #0c5460; }",
            ".metadata { background: #f8f9fa; padding: 10px; margin-bottom: 20px; }",
            "</style>",
            "</head>",
            "<body>",
            "<h1>Layout Extraction Visualization</h1>"
        ]
        
        # Metadata
        metadata = layout_data.get("metadata", {})
        html_lines.append("<div class='metadata'>")
        html_lines.append(f"<strong>File:</strong> {metadata.get('filename', 'unknown')}<br>")
        html_lines.append(f"<strong>Pages:</strong> {layout_data.get('total_pages', 0)}")
        html_lines.append("</div>")
        
        # Pages
        for page in layout_data.get("pages", []):
            page_num = page.get("page", 0)
            blocks = page.get("blocks", [])
            
            html_lines.append(f"<div class='page'>")
            html_lines.append(f"<h2>Page {page_num}</h2>")
            
            for block in blocks:
                text = block.get("text", "").replace("<", "&lt;").replace(">", "&gt;")
                
                css_class = "block"
                if block.get("is_potential_heading"):
                    css_class += " heading"
                elif block.get("list_type"):
                    css_class += " list"
                
                indent_px = block.get("indent_level", 0) * 20
                
                html_lines.append(f"<div class='{css_class}' style='margin-left: {indent_px}px;'>")
                html_lines.append(f"{text}")
                
                # Metadata badges
                badges = []
                if block.get("list_type"):
                    badges.append(f"<span style='background:#6c757d;color:white;padding:2px 5px;border-radius:3px;font-size:10px;'>{block['list_type']}</span>")
                if block.get("font"):
                    badges.append(f"<span style='background:#6c757d;color:white;padding:2px 5px;border-radius:3px;font-size:10px;'>{block['font']} {block.get('size', 0)}pt</span>")
                
                if badges:
                    html_lines.append("<br>" + " ".join(badges))
                
                html_lines.append("</div>")
            
            html_lines.append("</div>")
        
        html_lines.append("</body></html>")
        
        # Save
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write('\n'.join(html_lines))
        
        logger.info(f"✓ Layout visualization saved to: {output_file}")


if __name__ == "__main__":
    import sys
    import json
    
    logging.basicConfig(
        level=logging.INFO,
        format='[%(levelname)s] %(message)s'
    )
    
    if len(sys.argv) < 3:
        print("Usage: python document_reconstructor.py ai_output.json output.docx")
        sys.exit(1)
    
    input_json = sys.argv[1]
    output_docx = sys.argv[2]
    
    # Load AI output
    with open(input_json, 'r', encoding='utf-8') as f:
        ai_output = json.load(f)
    
    # Reconstruct
    reconstruct_document(ai_output, output_docx, generate_report=True)
    
    print(f"✓ Document reconstructed: {output_docx}")
